#!/usr/bin/env python3
"""Generate the Word document for the ImageNodeEditor project."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Helper: set font ──
def set_font(style, font_name, size=Pt(11), bold=False):
    style.font.name = font_name
    style.font.size = size
    style.font.bold = bold
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = style.element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

# ── Styles ──
set_font(doc.styles['Normal'], '宋体', Pt(12))
set_font(doc.styles['Heading 1'], '宋体', Pt(16), bold=True)
set_font(doc.styles['Heading 2'], '宋体', Pt(14), bold=True)
set_font(doc.styles['Heading 3'], '宋体', Pt(12), bold=True)
set_font(doc.styles['List Bullet'], '宋体', Pt(12))

def set_run_font(run):
    run.font.name = '宋体'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')

def set_paragraph_font(p):
    for run in p.runs:
        set_run_font(run)

def set_cell_font(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')
    run.bold = bold

# ── Title ──
title = doc.add_heading('图像节点编辑器 — 课程大作业', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '宋体'
    run.font.size = Pt(22)
    run.font.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')

p = doc.add_paragraph('ImageNodeEditor — 基于节点的图像处理工作流工具')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    set_run_font(run)
    run.font.italic = True
    run.font.size = Pt(12)

doc.add_paragraph('')

# ═══════════════════════════════════════════
# 1. 概述
# ═══════════════════════════════════════════
doc.add_heading('一、概述', level=1)

doc.add_heading('1.1 项目背景', level=2)
p = doc.add_paragraph(
    '本课程大作业实现了一个基于节点编辑器的图像处理工作流工具（ImageNodeEditor），'
    '用户通过拖拽连接节点的方式构建图像处理管线，无需编写代码即可完成复杂的图像处理任务。'
    '项目使用 C++17 和 Qt5 框架开发，采用 CMake 构建系统，支持 Windows 平台。'
)
set_paragraph_font(p)

doc.add_heading('1.2 功能简介', level=2)
p = doc.add_paragraph(
    'ImageNodeEditor 提供了 33 种不同类型的图像处理节点，分为输入、输出、滤波、'
    '风格化、色彩调整、几何变换、转换、多端口等八大类别。用户可以通过可视化界面拖拽添加节点、'
    '连线建立数据通路、调节参数实时预览效果，并支持工作流的保存与加载。'
)
set_paragraph_font(p)

doc.add_heading('1.3 技术栈', level=2)
items = [
    ('编程语言', 'C++17'),
    ('GUI 框架', 'Qt 5.15+（Core, Gui, Widgets）'),
    ('构建工具', 'CMake 3.16+, Visual Studio 2022'),
    ('设计模式', '单例模式（NodeRegistry）、工厂方法模式、观察者模式（Signal/Slot）'),
    ('数据结构', '有向无环图（DAG）、拓扑排序（Kahn 算法）、LUT 预计算'),
]
for label, value in items:
    p = doc.add_paragraph(style='List Bullet')
    runner = p.add_run(f'{label}：')
    runner.bold = True
    p.add_run(value)
    set_paragraph_font(p)

# ═══════════════════════════════════════════
# 2. 结构设计
# ═══════════════════════════════════════════
doc.add_heading('二、结构设计', level=1)

doc.add_heading('2.1 整体架构', level=2)
p = doc.add_paragraph('项目采用分层模块化设计，主要分为以下七个模块：')
set_paragraph_font(p)

# Architecture table
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Shading Accent 1'
headers = ['模块', '目录', '职责']
for i, h in enumerate(headers):
    set_cell_font(table.rows[0].cells[i], h, bold=True)

data = [
    ('Core 核心框架', 'main/core/', 'Node 基类、端口定义、数据类型、节点注册工厂'),
    ('Algorithms 算法层', 'main/algorithms/', '无状态纯函数图像处理算法，线程安全'),
    ('Nodes 节点实现', 'main/nodes/', '33 个具体节点，继承 Node 基类实现 process()'),
    ('Engine 工作流引擎', 'main/engine/', 'WorkflowEngine 管理 DAG 节点拓扑与执行'),
    ('GUI 图形界面', 'main/gui/', '场景、视图、节点图形项、属性面板、日志面板等'),
    ('IO 序列化', 'main/io/', 'WorkflowSerializer JSON 序列化/反序列化'),
    ('CLI 命令行接口', 'main/cli/', 'CommandLineRunner 无 GUI 执行工作流'),
]
for i, (mod, path, desc) in enumerate(data):
    set_cell_font(table.rows[i + 1].cells[0], mod)
    set_cell_font(table.rows[i + 1].cells[1], path)
    set_cell_font(table.rows[i + 1].cells[2], desc)

doc.add_paragraph('')

doc.add_heading('2.2 核心类设计', level=2)

classes = [
    ('Node 基类：', '所有处理节点的抽象基类，继承自 QObject。定义了端口（QVector<Port>）、'
     '参数系统（QMap<QString, QVariant> + ParamBound 边界约束）、'
     '纯虚函数 process() 和 clone()。参数变更时发射 paramsChanged() 信号。'),
    ('NodeRegistry 单例工厂：', '全局唯一的节点注册表，管理 33 种节点类型的创建函数。'
     '提供 categorizedEntries() 按预设分类顺序输出，支持左侧面板分类展示。'),
    ('WorkflowEngine 引擎：', '管理节点（QMap<QUuid, Node*>）和连线（QVector<Connection>）。'
     '核心功能包括：端口类型兼容性检查、DAG 环检测（Kahn 拓扑排序）、'
     '可达性分析（双向 BFS 找出完整通路）、按拓扑顺序依次执行。'),
    ('DataPacket 数据包：', '节点间传递的数据载体，支持单张 QImage 或 QVector<QImage> 列表。'
     '提供 isValid()、image()、imageList()、firstImage() 等便捷接口。'),
]
for name, desc in classes:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(name)
    r.bold = True
    p.add_run(desc)
    set_paragraph_font(p)

doc.add_heading('2.3 数据流架构', level=2)
p = doc.add_paragraph(
    '用户通过 GUI 拖拽建立节点连线，形成有向无环图（DAG）。执行时，WorkflowEngine 首先进行'
    '双向 BFS 可达性分析，从起点（入度为 0 且有输出的节点）正向可达集和终点（出度为 0 且有输入的节点）'
    '反向可达集的交集确定"活跃节点集"，再进行拓扑排序，按序执行每个节点的 process() 方法，'
    '将输出 DataPacket 存入结果映射表供下游使用。'
)
set_paragraph_font(p)

# ═══════════════════════════════════════════
# 3. 具体功能实现
# ═══════════════════════════════════════════
doc.add_heading('三、具体功能实现', level=1)

doc.add_heading('3.1 节点系统', level=2)
p = doc.add_paragraph(
    '33 个处理节点按功能分为 8 类。每个节点继承 Node 基类，在构造函数中声明端口和默认参数，'
    '实现 process() 方法调用 ImageAlgorithm 中的对应算法函数。节点支持参数边界约束'
    '（setParamBound），属性面板自动为数值参数生成带范围限制的 QSpinBox/QDoubleSpinBox。'
)
set_paragraph_font(p)

# Node table
table2 = doc.add_table(rows=9, cols=2)
table2.style = 'Light Shading Accent 1'
set_cell_font(table2.rows[0].cells[0], '分类', bold=True)
set_cell_font(table2.rows[0].cells[1], '包含节点', bold=True)

node_data = [
    ('输入', '图像输入'),
    ('输出', '图像输出、图像查看器'),
    ('滤波', '高斯模糊、锐化、边缘检测、清晰度'),
    ('风格化', '怀旧、像素化、暗角、铅笔画、卡通风格'),
    ('色彩调整', '亮度/对比度、饱和度、色相偏移、伽马校正、反色、曝光度、'
                 '自然饱和度、白平衡、阴影/高光、黑/白场、色调曲线'),
    ('几何变换', '裁剪、缩放、旋转、水印'),
    ('转换', '灰度化、阈值、冷暖色调、褪色'),
    ('多端口', '通道分离、通道合并'),
]
for i, (cat, nodes) in enumerate(node_data):
    set_cell_font(table2.rows[i + 1].cells[0], cat)
    set_cell_font(table2.rows[i + 1].cells[1], nodes)

doc.add_paragraph('')

doc.add_heading('3.2 工作流引擎', level=2)
p = doc.add_paragraph('WorkflowEngine 是整个应用的核心控制模块，负责管理节点生命周期、连线拓扑和执行调度：')
set_paragraph_font(p)

items = [
    '拓扑排序与环检测：使用 Kahn 算法计算节点的线性执行顺序，若结果节点数少于总节点数则判定存在环。',
    '端口兼容性检查：连线时检查数据类型是否匹配（ColorImage ↔ GrayImage ↔ Any），目标端口是否已被占用，防止自连。',
    '可达性分析：执行前通过正向 BFS（从起点）和反向 BFS（从终点）取交集，确定活跃节点集，跳过孤立节点。',
    '替换节点：右键菜单触发，创建新节点后尝试重连所有原始连线，若端口不兼容则自动回滚，保证工作流完整性。',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    set_paragraph_font(p)

doc.add_heading('3.3 图像处理算法', level=2)
p = doc.add_paragraph('ImageAlgorithm 提供 20+ 种无状态纯函数图像处理算法，所有函数均线程安全。主要实现：')
set_paragraph_font(p)

items = [
    '色彩调整：HSL 色彩空间转换实现饱和度调节和色相偏移；LUT 预计算加速亮度/对比度、伽马校正、曝光度调整。',
    '滤波算法：可分离高斯模糊（水平+垂直一维卷积）；Sobel 边缘检测（3×3 梯度算子 + 双阈值）；4-邻域拉普拉斯锐化。',
    '风格化效果：像素化（块采样 + 近邻填充）；暗角（径向余弦衰减）；铅笔画（反色 + 高斯模糊 + 色彩减淡）；卡通（颜色量化 + Sobel 边缘描边）。',
    '色调曲线：白/R/G/B 四通道独立 256 级 LUT，支持 Catmull-Rom 平滑曲线插值，交互式编辑器可拖拽添加/调整节点。',
    '高级色彩：白平衡（色温/色调增益 + 灰度世界自动校正）；自然饱和度（智能增强低饱和区域，保护肤色）；阴影/高光（基于亮度掩膜的独立调节）。',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    set_paragraph_font(p)

doc.add_heading('3.4 图形界面', level=2)
p = doc.add_paragraph('GUI 模块基于 Qt5 Graphics View 框架实现，采用 MVVM 风格架构：')
set_paragraph_font(p)

items = [
    'NodeScene（QGraphicsScene）：管理所有图形项，处理鼠标事件（拖拽连线、选择、框选），绘制网格背景。',
    'NodeView（QGraphicsView）：支持滚轮缩放（0.2× ~ 5×）、中键/右键拖拽平移画布。',
    'NodeGraphicsItem：自定义 QGraphicsItem，绘制圆角矩形节点，标题栏使用分类颜色标识，输入/输出端口分布在左右两侧。',
    'ConnectionGraphicsItem：三次贝塞尔曲线连接线，双击节点打开预览对话框，右键可删除连线。',
    'PropertyPanel：自动根据节点参数类型生成编辑控件（QSpinBox、QDoubleSpinBox、QCheckBox、QLineEdit），文件路径附带浏览对话框。',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    set_paragraph_font(p)

doc.add_heading('3.5 序列化与命令行', level=2)
p = doc.add_paragraph(
    'WorkflowSerializer 将工作流保存为 JSON 格式，包含节点（类型、ID、位置、参数）'
    '和连线（源/目标节点 ID 及端口索引）信息。支持从命令行无 GUI 执行工作流'
    '（ImageNodeEditor.exe --no-gui --workflow xxx.json），'
    '并可覆盖输入/输出图像路径（--image / --output 参数）。'
)
set_paragraph_font(p)

# ═══════════════════════════════════════════
# 4. 反思
# ═══════════════════════════════════════════
doc.add_heading('四、反思', level=1)

doc.add_heading('4.1 设计亮点', level=2)
items = [
    '模块化设计：算法层与节点层完全解耦，ImageAlgorithm 的纯函数设计便于单元测试和复用。',
    '参数边界约束：ParamBound 机制确保 UI 控件自动限制在有效范围，防止越界错误。',
    '节点替换回滚：在处理新节点端口不兼容时能自动撤销所有变更，保持工作流一致性。',
    '可达性分析：执行前的双向 BFS 分析避免执行孤立节点，提高效率并给出清晰错误提示。',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    set_paragraph_font(p)

doc.add_heading('4.2 存在问题与改进方向', level=2)
items = [
    '撤销/重做功能缺失：当前删除节点或连线后无法恢复，替换节点的回滚机制也未覆盖正常删除操作。',
    '类别名称不一致：33 个节点的 category() 返回值中英文混用，导致早期版本的颜色渲染异常（已修复）。',
    '锐化算法精度：当前使用 4-邻域拉普拉斯核，效果弱于标准 8-邻域或 Unsharp Mask 方案。',
    '缺失图像格式支持：WebP、SVG 等格式依赖 Qt 插件，conda 环境下的 imageformats 插件部署需手动配置。',
    '实时预览缺失：当前需手动执行工作流才能看到结果，缺少参数调整时的实时输出预览。',
    '缺乏性能优化：大图像处理时缺少进度反馈和异步执行机制（QThread/QFuture），界面可能短暂无响应。',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    set_paragraph_font(p)

doc.add_heading('4.3 经验总结', level=2)
p = doc.add_paragraph(
    '通过本项目深入实践了 C++17 和 Qt5 GUI 编程，掌握了 Graphics View 框架的自定义图形项和事件处理机制、'
    '基于 DAG 的工作流引擎设计、以及多种图像处理算法的实现。项目从最初的简单滤镜工具逐步演进为功能较完整的'
    '节点编辑器，过程中进行了多次重构（算法与节点解耦、参数系统统一化、引擎执行优化），'
    '验证了良好的架构设计对项目可维护性和可扩展性的重要性。'
)
set_paragraph_font(p)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), '大作业报告_ImageNodeEditor.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')

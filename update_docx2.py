"""Update the Word report with auto-connect and undo fixes."""

import docx

path = r"C:\Users\sunli\Desktop\期末大作业\main\大作业报告_ImageNodeEditor.docx"
doc = docx.Document(path)

changes = []

def set_para_text(para, new_text):
    if para.runs:
        first = para.runs[0]
        for run in para.runs[1:]:
            run._element.getparent().remove(run._element)
        first.text = new_text
    else:
        para.text = new_text

# P34 - auto-connect description
p34 = doc.paragraphs[34]
new_p34 = (
    "替换节点：右键菜单触发，创建新节点后尝试重连所有原始连线，若端口不兼容则自动回滚，保证工作流完整性；"
    "F6 一键自动连线，使用可用输出池算法，支持多端口节点（ChannelSplit 三路扇出、ChannelMerge 三路汇聚），"
    "类型兼容匹配，未获输入的节点不向池中贡献输出，Any 类型输出消费后即移除防止复用。"
)
if p34.text != new_p34:
    set_para_text(p34, new_p34)
    changes.append("P34: Updated auto-connect description")

# P58 - undo reflection
p58 = doc.paragraphs[58]
new_p58 = (
    "撤销/重做功能不完善：已实现 Ctrl+Z 撤销最多 50 步（添加/删除节点/连线、拖拽移动位置），"
    "修复了撤销恢复时快照重复入栈导致无限循环的 Bug，以及节点位置未持久化导致恢复后全部重叠的问题。"
    "但缺少重做(Redo)功能，替换节点回滚机制在极端场景下可能残留不一致状态。"
)
if p58.text != new_p58:
    set_para_text(p58, new_p58)
    changes.append("P58: Updated undo reflection")

doc.save(path)
print("=== Changes Made ===")
for c in changes:
    print(f"  ✓ {c}")
